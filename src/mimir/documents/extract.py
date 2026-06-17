"""Text extraction for document ingestion (DESIGN §8).

Resolves an extractor by file extension and returns a list of ``ExtractedUnit`` — a span of
text plus a **locator** (page, section, …) that becomes provenance downstream. Plain text and
markdown are handled in core with zero dependencies; PDF lives behind the optional
``[documents]`` extra so the runtime contract (Python + SQLite, nothing else) holds for anyone
who doesn't need it.

A missing optional extractor fails loud with an install instruction — never a silent skip.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path

from ..errors import IngestError

log = logging.getLogger("mimir.documents")

_TEXT_EXTS = {".txt", ".text", ""}
_MARKDOWN_EXTS = {".md", ".markdown", ".mdown"}
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}

# ATX markdown heading, e.g. "## Section title".
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")


@dataclass(slots=True)
class ExtractedUnit:
    """A span of text with a human-readable locator (becomes part of provenance)."""

    text: str
    locator: str  # e.g. "p.3", a section heading, or "" for an un-located whole file


def extract(path: Path) -> list[ExtractedUnit]:
    """Extract a document into located units. Raises ``IngestError`` on unsupported types."""
    ext = path.suffix.lower()
    if ext in _TEXT_EXTS:
        return _extract_text(path)
    if ext in _MARKDOWN_EXTS:
        return _extract_markdown(path)
    if ext in _PDF_EXTS:
        return _extract_pdf(path)
    if ext in _DOCX_EXTS:
        return _extract_docx(path)
    raise IngestError(
        f"unsupported document type {ext!r} for {path.name}. Supported: .txt, .md in core; "
        f".pdf and .docx via the optional extra (pip install 'mimir-0[documents]')."
    )


def _read_text(path: Path) -> str:
    # Documents are best-effort text; undecodable bytes are replaced rather than crashing the
    # ingest. This is a convenience layer, not core memory integrity.
    raw = path.read_bytes()
    text = raw.decode("utf-8", errors="replace")
    if "�" in text:
        log.warning("extract: %s had non-UTF-8 bytes; replaced during decode", path.name)
    return text


def _extract_text(path: Path) -> list[ExtractedUnit]:
    text = _read_text(path)
    return [ExtractedUnit(text=text, locator="")] if text.strip() else []


def _extract_markdown(path: Path) -> list[ExtractedUnit]:
    """Split markdown into sections by ATX headings; the heading becomes the locator."""
    text = _read_text(path)
    units: list[ExtractedUnit] = []
    current_heading = ""
    buffer: list[str] = []

    def flush() -> None:
        body = "\n".join(buffer).strip()
        if body:
            units.append(ExtractedUnit(text=body, locator=current_heading))

    for line in text.splitlines():
        match = _HEADING_RE.match(line)
        if match:
            flush()
            current_heading = match.group(2).strip()
            buffer = [line]  # keep the heading line in the section body for context
        else:
            buffer.append(line)
    flush()
    return units


def _extract_pdf(path: Path) -> list[ExtractedUnit]:
    try:
        import pypdf
    except ImportError as exc:
        raise IngestError(
            "PDF ingestion needs the optional extra. Install it with: "
            "pip install 'mimir-0[documents]'"
        ) from exc

    reader = pypdf.PdfReader(str(path))
    units: list[ExtractedUnit] = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            units.append(ExtractedUnit(text=page_text, locator=f"p.{i + 1}"))
    if not units:
        log.warning("extract: %s yielded no extractable text (scanned/image PDF?)", path.name)
    return units


def _docx_block_items(parent: object) -> list[object]:
    """Paragraphs and tables of a docx body (or table cell) in **document order**.

    python-docx exposes ``.paragraphs`` and ``.tables`` as separate, unordered collections —
    iterating only paragraphs silently drops every table (and a table-structured document, like a
    safety matrix, can be almost entirely tables). Walking the underlying XML keeps the two
    interleaved, so table content stays attached to the heading that precedes it."""
    from docx.document import Document as _DocumentObject
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _DocumentObject):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        return []
    items: list[object] = []
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            items.append(Paragraph(child, parent))
        elif isinstance(child, CT_Tbl):
            items.append(Table(child, parent))
    return items


def _docx_table_lines(table: object, seen: set[int]) -> list[str]:
    """Render a table as readable rows — non-empty cells joined by ' | ', one line per row. Merged
    cells (which python-docx repeats per grid position) are de-duped via ``seen``; nested tables in
    a cell are recursed into."""
    from docx.table import Table

    lines: list[str] = []
    for row in table.rows:  # type: ignore[attr-defined]
        cells: list[str] = []
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen:  # a merged cell already emitted (horizontal or vertical span)
                continue
            seen.add(tc_id)
            parts = [(cell.text or "").strip()]
            for block in _docx_block_items(cell):
                if isinstance(block, Table):
                    parts.extend(_docx_table_lines(block, seen))
            text = " ".join(p for p in parts if p).strip()
            if text:
                cells.append(text)
        if cells:
            lines.append(" | ".join(cells))
    return lines


def _extract_docx(path: Path) -> list[ExtractedUnit]:
    """Split a .docx into sections by Word heading styles (Heading 1/2/…, Title); the heading text
    becomes the locator, like markdown. Walks paragraphs **and tables** in document order, so
    table-structured documents aren't lost. (python-docx reads .docx only, not legacy .doc.)"""
    try:
        import docx  # noqa: F401  (presence check; the walker imports submodules lazily)
    except ImportError as exc:
        raise IngestError(
            "DOCX ingestion needs the optional extra. Install it with: "
            "pip install 'mimir-0[documents]'"
        ) from exc
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(str(path))
    units: list[ExtractedUnit] = []
    current_heading = ""
    buffer: list[str] = []

    def flush() -> None:
        body = "\n".join(buffer).strip()
        if body:
            units.append(ExtractedUnit(text=body, locator=current_heading))

    for block in _docx_block_items(document):
        if isinstance(block, Paragraph):
            text = (block.text or "").strip()
            if not text:
                continue
            style = (block.style.name if block.style else "") or ""
            if style.startswith("Heading") or style == "Title":
                flush()
                current_heading = text
                buffer = [text]  # keep the heading line in the section body for context
            else:
                buffer.append(text)
        elif isinstance(block, Table):
            rows = _docx_table_lines(block, set())
            if rows:
                buffer.append("\n".join(rows))
    flush()
    if not units:
        log.warning("extract: %s yielded no extractable text", path.name)
    return units
